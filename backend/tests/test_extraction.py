"""Tests for the extraction stage's per-format extractors.

Each extractor is checked on the two things the pipeline downstream depends on:
the text that comes out is the text that went in, and it arrives with the page
structure the *format* actually has. Only a PDF has real pages; the rest must
collapse to exactly one, because inventing page breaks would attribute chunks to
pages the source never had.
"""

from collections.abc import Callable

import pytest

from services.extraction import (
    DocxExtractor,
    HtmlExtractor,
    PdfExtractor,
    TextExtractor,
    extractor_for,
    single_page,
)
from dtos.responses import DocType


class TestExtractorRegistry:
    def test_every_ingestible_type_has_an_extractor(self) -> None:
        ingestible = [
            doc_type for doc_type in DocType if doc_type is not DocType.unknown
        ]

        assert all(extractor_for(doc_type) is not None for doc_type in ingestible)

    def test_unknown_has_no_extractor(self) -> None:
        assert extractor_for(DocType.unknown) is None


class TestSinglePage:
    """The paginationless degradation: many blocks in, exactly one page out."""

    def test_joins_blocks_into_one_page_separated_by_a_blank_line(self) -> None:
        assert single_page(["first", "second"]) == ["first\n\nsecond"]

    def test_drops_blank_blocks_and_collapses_internal_whitespace(self) -> None:
        assert single_page(["  spaced   out ", "   ", "\t", "next"]) == [
            "spaced out\n\nnext"
        ]

    def test_no_blocks_still_yields_one_page(self) -> None:
        assert single_page([]) == [""]


class TestPdfExtractor:
    def test_extracts_one_entry_per_page(
        self, make_pdf: Callable[[list[str]], bytes]
    ) -> None:
        pages = PdfExtractor().extract(make_pdf(["Hello page one", "Second page"]))

        assert len(pages) == 2
        assert "Hello page one" in pages[0]
        assert "Second page" in pages[1]

    def test_rejects_non_pdf_bytes(self) -> None:
        with pytest.raises(ValueError, match="not a readable PDF"):
            PdfExtractor().extract(b"this is plainly not a pdf")


class TestDocxExtractor:
    def test_extracts_paragraphs_onto_a_single_page(
        self, make_docx: Callable[[list[str]], bytes]
    ) -> None:
        pages = DocxExtractor().extract(make_docx(["First para.", "Second para."]))

        assert pages == ["First para.\n\nSecond para."]

    def test_keeps_table_text_in_document_order(self) -> None:
        from docx import Document
        from io import BytesIO

        document = Document()
        document.add_paragraph("Before.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Left"
        table.rows[0].cells[1].text = "Right"
        document.add_paragraph("After.")
        buffer = BytesIO()
        document.save(buffer)

        # A table is body content, not an appendix: reading paragraphs alone
        # would drop it, and appending it would put it after "After.".
        assert DocxExtractor().extract(buffer.getvalue()) == [
            "Before.\n\nLeft Right\n\nAfter."
        ]

    def test_rejects_bytes_that_are_not_a_docx(self) -> None:
        with pytest.raises(ValueError, match="not a readable DOCX"):
            DocxExtractor().extract(b"not a word document")


class TestHtmlExtractor:
    def test_extracts_readable_text_onto_a_single_page(self) -> None:
        pages = HtmlExtractor().extract(
            b"<html><body><p>One</p><p>Two</p></body></html>"
        )

        assert pages == ["One\n\nTwo"]

    def test_drops_script_style_and_head_content(self) -> None:
        html = (
            b"<html><head><title>Tab title</title><style>p{color:red}</style></head>"
            b"<body><script>var tracker = 1;</script><p>Real prose.</p></body></html>"
        )

        # Script and style text is never shown to a reader; embedding it would
        # put minified JavaScript into the chunks.
        assert HtmlExtractor().extract(html) == ["Real prose."]

    def test_block_tags_break_paragraphs_but_inline_tags_do_not(self) -> None:
        html = b"<h1>Heading</h1><p>A <b>bold</b>face word.</p>"

        # <b> sits inside a word, so splitting on it would invent "bold" and
        # "face" as separate tokens.
        assert HtmlExtractor().extract(html) == ["Heading\n\nA boldface word."]

    def test_markup_free_bytes_are_still_readable_text(self) -> None:
        assert HtmlExtractor().extract(b"just words") == ["just words"]


class TestTextExtractor:
    def test_decodes_onto_a_single_page(self) -> None:
        assert TextExtractor().extract(b"Line one.\n\nLine two.") == [
            "Line one.\n\nLine two."
        ]

    def test_strips_a_byte_order_mark(self) -> None:
        assert TextExtractor().extract("\ufeffLeading text".encode()) == [
            "Leading text"
        ]

    def test_undecodable_bytes_are_replaced_rather_than_rejected(self) -> None:
        # A file that is mostly prose should still be ingestible.
        pages = TextExtractor().extract(b"good text \xff\xfe more text")

        assert pages[0].startswith("good text")
        assert pages[0].endswith("more text")
